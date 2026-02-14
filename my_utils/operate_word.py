from docx import Document
import os, shutil, tempfile, copy
from my_utils.operate_excel import Pair
from my_utils.utils import get_all_image_paths
from docx.shared import Inches
from pathlib import Path
from PIL import Image, ImageOps
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Mm
from docxcompose.composer import Composer
from docx.enum.section import WD_SECTION

def is_word_empty(doc):
    # 判断段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            return False

    # 判断表格文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    return False

    # 判断图片（内嵌图片）
    if len(doc.inline_shapes) > 0:
        return False

    return True

def copy_page_setup(src_sec, dst_sec):
    # 页边距
    dst_sec.top_margin = src_sec.top_margin
    dst_sec.bottom_margin = src_sec.bottom_margin
    dst_sec.left_margin = src_sec.left_margin
    dst_sec.right_margin = src_sec.right_margin
    dst_sec.gutter = src_sec.gutter

    # 纸张与方向
    dst_sec.page_width = src_sec.page_width
    dst_sec.page_height = src_sec.page_height
    dst_sec.orientation = src_sec.orientation

    # 页眉页脚距离
    dst_sec.header_distance = src_sec.header_distance
    dst_sec.footer_distance = src_sec.footer_distance

def concat_two_words(doc1, doc2, mode, out_path=None):
    if isinstance(doc1, str):
        d1_temp = Document(doc1)
    else:
        d1_temp = doc1
    if not mode:
        d1 = copy.deepcopy(d1_temp)
    else:
        d1 = d1_temp
    d2 = Document(doc2) if isinstance(doc2, str) else doc2

    # 1) 先在 d1 末尾插入“分节符(下一页)”
    new_sec = d1.add_section(WD_SECTION.NEW_PAGE)

    # 2) 把 d2 的第一页所属 section 的页面设置拷到这个新节上
    copy_page_setup(d2.sections[0], new_sec)

    # 3) 再用 docxcompose 追加内容（图片/样式/编号等更稳）
    composer = Composer(d1)
    composer.append(d2)

    if out_path:
        composer.save(out_path)
    return composer.doc

def append_fullpage_image_center(doc, pics_floader='../配置/pics', margin_mm=5):
    pics_paths = get_all_image_paths(pics_floader)

    with tempfile.TemporaryDirectory() as tmpdir:
        for idx, pic in enumerate(pics_paths):
            # 新建一个新页 section
            sec = doc.add_section(WD_SECTION_START.NEW_PAGE)

            # A4
            sec.page_width = Mm(210)
            sec.page_height = Mm(297)

            # 边距（打印建议 5~10mm）
            sec.left_margin = Mm(margin_mm)
            sec.right_margin = Mm(margin_mm)
            sec.top_margin = Mm(margin_mm)
            sec.bottom_margin = Mm(margin_mm)

            # 版心尺寸
            max_w = sec.page_width - sec.left_margin - sec.right_margin
            max_h = sec.page_height - sec.top_margin - sec.bottom_margin

            # 读取图片 + 纠正EXIF旋转
            with Image.open(pic) as im:
                im = ImageOps.exif_transpose(im)  # 关键：修正90度旋转
                w_px, h_px = im.size
                fixed_path = os.path.join(tmpdir, f"fixed_{idx}.png")
                im.save(fixed_path, format="PNG")  # 用PNG避免再次受EXIF影响

            ratio = w_px / h_px

            # 等比最大
            if (max_w / max_h) >= ratio:
                target_h = int(max_h)
                target_w = int(target_h * ratio)
            else:
                target_w = int(max_w)
                target_h = int(target_w / ratio)

            # 段落水平居中
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0

            run = p.add_run()
            run.add_picture(fixed_path, width=target_w)
    return doc

#提取照片
def extract_images_from_docx(docx_path, output_folder):
    # 创建输出文件夹，如果不存在的话
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 打开docx文件
    doc = Document(docx_path)

    # 初始化图片计数器
    image_count = 0

    # 遍历文档中的所有段落和图形
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            # 获取图片数据
            image = rel.target_part.blob
            # 获取图片格式
            extension = rel.target_part.content_type.split('/')[-1]
            # 保存图片
            image_filename = os.path.join(output_folder, f'image_{image_count}.{extension}')
            with open(image_filename, 'wb') as img_file:
                img_file.write(image)

    # print(f'提取完成，共提取了 {image_count} 张图片。')

def copy_template(mode, base_path, name_concat, count, sub = None, in_floader=False, template_sub_path = []):
    template_name = ['两个内地人模版', '一港一内地模版']
    if sub != None:
        docx_name = name_concat + '_' + mode + '_' + sub + '.docx'
    else:
        docx_name = name_concat + '_' + mode + '.docx'
    name = template_name[int(bool(count))]
    if sub == None:
        kaidan_template_path = Path(f'模版/{mode}/{name}.docx')
    else:
        kaidan_template_path = Path(f'模版/{mode}/{name}/{sub}.docx')
    if len(template_sub_path) != 0:
        kaidan_template_path_temp = kaidan_template_path.parent / Path(*template_sub_path) / kaidan_template_path.name
        if os.path.exists(kaidan_template_path_temp):
            kaidan_template_path = kaidan_template_path_temp
    if in_floader:
        kaidan_base = os.path.join(base_path, mode, name_concat)
    else:
        kaidan_base = os.path.join(base_path, mode)
    os.makedirs(kaidan_base, exist_ok=True)
    kaidan_path = os.path.join(kaidan_base, docx_name)
    if os.path.exists(kaidan_path):
        os.remove(kaidan_path)
    shutil.copy(kaidan_template_path, kaidan_path)
    return kaidan_path

def number_to_chinese(s):
    s_duizhao = {
        '0':"零",'1':"壹",'2':"贰",'3':"叁",'4':"肆",
        '5':"伍",'6':"陆",'7':"柒",'8':"捌",'9':"玖"
    }
    #小数点前最大到千亿 可继续增加
    if len(s) == 1:
        return s_duizhao[s] + '元'
    else:
        s_danwei = {
            0:'',1:'拾',2:'佰',3:'仟',4:'万',8:'亿'
        }
        re_re = ""
        if '.' in s :
            s_1 = s[0:s.find('.')]
            s_2 = s[s.find('.')+1:]
            flag_1 = 0
            s_1 = s_1[::-1]

            for i in s_1:
                if flag_1 != 0 and flag_1 % 4 == 0:
                    re_re += s_danwei[flag_1]
                if i != '0':
                    re_re += s_danwei[flag_1 % 4]
                re_re += s_duizhao[i]
                flag_1 += 1
            re_re = re_re[::-1]
            while re_re.endswith('零') and len(re_re)>1:
                re_re = re_re[:-1]
            # 中文数字没有连续的零
            while "零零" in re_re:
                re_re = re_re.replace("零零", "零")
            re_re += '点'
            for i in s_2:
                re_re += s_duizhao[i]
        else:
            s_1 = s
            s_1 = s_1[::-1]
            flag_1 = 0

            for i in s_1:
                if flag_1 != 0 and flag_1 % 4 == 0:
                    re_re += s_danwei[flag_1]
                if i != '0':
                    re_re += s_danwei[flag_1 % 4]
                re_re += s_duizhao[i]
                flag_1 += 1
            re_re = re_re[::-1]
            while re_re.endswith('零') and len(re_re)>1:
                re_re = re_re[:-1]
            # 中文数字没有连续的零
            while "零零" in re_re:
                re_re = re_re.replace("零零", "零")

        if re_re[1] == '十' and re_re[0] == '一':
            re_re = re_re[1:]
    return re_re + '元'

def count_placeholder_occurrences(file_path, placeholder):
    count = 0
    doc = Document(file_path)

    def count_in_runs(runs):
        nonlocal count
        for run in runs:
            # run.text 中可能出现多次 placeholder，使用 str.count 统计
            count += run.text.count(placeholder)

    # 统计段落中的占位符
    for paragraph in doc.paragraphs:
        count_in_runs(paragraph.runs)

    # 统计表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count_in_runs(paragraph.runs)

    return count

def replace_text_with_same_format(file_path, placeholder, replacements):
    replacement_index = 0
    all_count = 0
    doc = Document(file_path)

    def replace_in_runs(runs):
        nonlocal all_count
        nonlocal replacement_index
        if not isinstance(replacements, str):
            for run in runs:
                while placeholder in run.text and replacement_index < len(replacements):
                    all_count += 1
                    # 找到占位符的位置
                    pos = run.text.find(placeholder)
                    # 分割文本
                    before = run.text[:pos]
                    after = run.text[pos + len(placeholder):]
                    # 替换文本，保持格式
                    run.text = before + replacements[replacement_index] + after
                    replacement_index += 1
        else:
            for run in runs:
                while placeholder in run.text:
                    all_count += 1
                    # 找到占位符的位置
                    pos = run.text.find(placeholder)
                    # 分割文本
                    before = run.text[:pos]
                    after = run.text[pos + len(placeholder):]
                    # 替换文本，保持格式
                    run.text = before + replacements + after

    # 处理文档中的每个段落
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph.runs)

    # 处理文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph.runs)
    return doc, all_count

def get_num_base(path, num=1):
    for i in range(num):
        path = os.path.dirname(path)
    return path

def get_image_paths(pairs:Pair, kaidan_path, in_folder_flag):
    pics = []
    base_path = get_num_base(kaidan_path, 2 + int(in_folder_flag))
    client_name = pairs.client.name
    entrusted_name = pairs.entrusted.name
    pics_temp = os.listdir(base_path)
    pics_temp = [i for i in pics_temp if not i.endswith('.info')]
    if '.DS_Store' in pics_temp:
        pics_temp.remove('.DS_Store')
    for k in pics_temp:
        if k.split('.')[0] == client_name:
            pics.append(k)
        elif k.split('.')[0] == entrusted_name:
            pics.append(k)
        if len(pics) == 2:
            break
    pics = [f"{client_name}.png", f"{client_name}反.png", f"{entrusted_name}.png", f"{entrusted_name}反.png"]
    return [os.path.join(base_path, pic) for pic in pics]

def replace_pic(doc, pairs, num, kaidan_path, position, num_spaces, scale = 1.0, in_folder_flag = True):
    image_paths = get_image_paths(pairs, kaidan_path, in_folder_flag)
    image_paths = image_paths[2*position:2*(position + 1)]
    # 计算目标段落的索引
    target_index = len(doc.paragraphs) - num
    
    # 确保文档至少有三个段落
    if target_index < 0:
        raise ValueError("文档段落不足三个，无法在倒数第三段插入图片。")
    
    # 在目标段落之后创建一个新段落用于插入图片
    paragraph = doc.paragraphs[target_index]
    new_paragraph = paragraph.insert_paragraph_before()  # 插入在目标段落之前
    
    # 在同一个段落内插入多个图片和空格
    for i, image_path in enumerate(image_paths):
        # 打开图片以获取其原始尺寸
        with Image.open(image_path) as img:
            original_width, original_height = img.size
            # 按比例计算新的宽度和高度
            scaled_width = original_width * scale
            scaled_height = original_height * scale
        
        # 转换尺寸到英寸（像素到英寸的转换因子通常是 96, 可根据实际情况调整）
        width_inches = scaled_width / 96
        height_inches = scaled_height / 96

        run = new_paragraph.add_run()
        # 插入图片，设置宽度和高度以保持比例
        run.add_picture(image_path, width=Inches(width_inches), height=Inches(height_inches))

        # 插入空格（除了最后一个图片之后）
        if i < len(image_paths) - 1:
            space_run = new_paragraph.add_run(' ' * num_spaces)
        
    # 将新段落移动到目标段落的后面
    paragraph._element.addnext(new_paragraph._element)
    return doc