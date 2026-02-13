from datetime import datetime, timedelta
import os, math, json, sys, base64, shutil, requests, cv2, copy, platform, tempfile
# import PIL
# import PIL.Image
# import PIL.ImageEnhance
from PyQt6.QtGui import QImage, QPixmap
from natsort import natsorted
import numpy as np
from pathlib import Path
from ruamel.yaml import YAML
from send2trash import send2trash
# from pypinyin import pinyin, Style
from PIL import Image, ImageOps
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Mm

def append_fullpage_image_center(doc, pics_floader='../files/pics', margin_mm=5):
    pics_paths = get_all_image_paths(pics_floader)

    with tempfile.TemporaryDirectory() as tmpdir:
        for idx, pic in enumerate(pics_paths):
            # 新建一个新页 section
            sec = doc.add_section(WD_SECTION_START.NEW_PAGE)

            # A4
            sec.page_width = Mm(210)
            sec.page_height = Mm(297)

            # 边距（打印建议 5~10mm）
            sec.left_margin = Mm(margin_mm)
            sec.right_margin = Mm(margin_mm)
            sec.top_margin = Mm(margin_mm)
            sec.bottom_margin = Mm(margin_mm)

            # 版心尺寸
            max_w = sec.page_width - sec.left_margin - sec.right_margin
            max_h = sec.page_height - sec.top_margin - sec.bottom_margin

            # 读取图片 + 纠正EXIF旋转
            with Image.open(pic) as im:
                im = ImageOps.exif_transpose(im)  # 关键：修正90度旋转
                w_px, h_px = im.size
                fixed_path = os.path.join(tmpdir, f"fixed_{idx}.png")
                im.save(fixed_path, format="PNG")  # 用PNG避免再次受EXIF影响

            ratio = w_px / h_px

            # 等比最大
            if (max_w / max_h) >= ratio:
                target_h = int(max_h)
                target_w = int(target_h * ratio)
            else:
                target_w = int(max_w)
                target_h = int(target_w / ratio)

            # 段落水平居中
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0

            run = p.add_run()
            run.add_picture(fixed_path, width=target_w)
    return doc

def get_all_image_paths(root_dir='../files/pics'):
    image_ext = ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp']
    image_paths = []

    for root, dirs, files in os.walk(root_dir):
        for file in files:
            ext = os.path.splitext(file)
            if ext[1].lower() in image_ext:
                try:
                    image_paths.append({os.path.join(root, file): int(ext[0])})
                except:
                    continue
    image_paths_temp = sorted(image_paths, key=lambda x: list(x.values())[0])
    image_paths_result = [next(iter(i.keys())) for i in image_paths_temp]

    return image_paths_result

def rgb_to_gray_with_three_channels(image):
    # 将 RGB 图像转换为灰度图像 (单通道)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # 将单通道灰度图像复制为 3 个通道
    gray_3_channels = cv2.merge([gray, gray, gray])
    return gray_3_channels

def get_config(config_path = '配置/conf.yaml'):
    # 初始化 YAML 处理器
    yaml = YAML()
    yaml.preserve_quotes = True  # 保留引号（如果 YAML 中有引号）
    with open(config_path, 'r', encoding='utf-8') as file:
        global_config = yaml.load(file)
    return global_config

def get_beijing_date():
    return datetime.now().strftime("%Y%m%d")

def write_config(conflg_yaml, save_path = '配置/conf.yaml'):
    # 初始化 YAML 处理器
    yaml = YAML()
    yaml.preserve_quotes = True  # 保留引号（如果 YAML 中有引号）
    with open(save_path, "w", encoding="utf-8") as f:
        yaml.dump(conflg_yaml, f)

#用下面这两个函数可以使用cv读取含有中文路径的文件
def cv_imread(file_path):
    # Read image as a byte array with np.fromfile
    image_array = np.fromfile(file_path, dtype=np.uint8)
    # Decode image using cv2.imdecode
    cv_img = cv2.imdecode(image_array, cv2.IMREAD_UNCHANGED)
    return cv_img

def find_in_catch_pic(name, floader, in_level=2):
    #从最新的开始索引
    catch_pic_paths = natsorted(os.listdir(floader), reverse=True)
    if '.DS_Store' in catch_pic_paths:
        catch_pic_paths.remove('.DS_Store')
    for i in catch_pic_paths:
        search = os.path.join(floader, i)
        cache_path = os.path.join(search, name)
        if os.path.exists(cache_path + '.info') and os.path.exists(cache_path + '.png') and os.path.exists(cache_path + '反.png'):
            return cache_path
    return ''

def delete_specific_files_and_folders(target_dir, del_dirname = None, del_filename = None):#, global_config = None):
    if del_dirname != None and del_filename != None:
        for root, dirs, files in os.walk(target_dir, topdown=False):
            # 删除 __pycache__ 文件夹
            if del_dirname != None:
                for dir_name in dirs:
                    if dir_name == del_dirname:
                        dir_path = os.path.join(root, dir_name)
                        shutil.rmtree(dir_path)  # 删除整个文件夹

            if del_filename != None:
                # 删除 .DS_Store 文件
                for file_name in files:
                    if file_name == del_filename:
                        file_path = os.path.join(root, file_name)
                        os.remove(file_path)  # 删除文件
    #将该文件夹隐藏
    path = './_internal'
    if os.path.exists(path):
        os.system(f'attrib +h "{path}"')
    # if global_config != None:
    #     path = './_internal'
    #     if global_config['del_packges_config']['del_flag'] and os.path.exists(path):
    #         packges = global_config['del_packges_config']['del_packges']
    #         if len(packges) != 0:
    #             for packge in packges:
    #                 path_t = os.path.join(path, packge)
    #                 if os.path.exists(path_t):
    #                     shutil.rmtree(path_t)

def get_internal_path(path:str):
    # 获取程序的资源文件路径
    if getattr(sys, 'frozen', False):  # 检查是否在打包后的模式下
        # 如果是打包后的程序，资源文件解压到 sys._MEIPASS 临时目录
        path_split = path.lstrip('./').rstrip('/')
        return os.path.join(sys._MEIPASS, path_split)
    else:
        return path

def check_catch_pic(catch_days):
    check_path = './模版/缓存照片'
    if catch_days != 0:
        day_pics = os.listdir(check_path)
        if '.DS_Store' in day_pics:
            day_pics.remove('.DS_Store')
        catch_days_floader = get_past_dates(catch_days)
        del_days = [i for i in day_pics if i not in catch_days_floader]
        if len(del_days) != 0:
            for i in del_days:
                send2trash(os.path.join(check_path, i))
        today_floader = os.path.join(check_path, catch_days_floader[0])
    else:
        catch_days_floader = get_past_dates(1)
        today_floader = os.path.join(check_path, catch_days_floader[0])
    os.makedirs(today_floader, exist_ok=True)
    return today_floader

def change_three_channel(cv_img):
    if cv_img is not None:
        if len(cv_img.shape) == 3:
            if cv_img.shape[2] == 4:  # 如果是 RGBA
                cv_temp = cv2.cvtColor(cv_img, cv2.COLOR_BGRA2BGR)
                return cv_temp
            elif cv_img.shape[2] == 1:  # 如果是单通道灰度图
                cv_temp = cv2.cvtColor(cv_img, cv2.COLOR_GRAY2BGR)
                return cv_temp
            else:
                return cv_img
        elif len(cv_img.shape) == 2:
            cv_temp = np.stack([cv_img] * 3, axis=-1)  # 在最后一维重复 3 次形成 RGB 格式
            return cv_temp
    return cv_img

def cv_imwrite(image, file_type, save_path):
    cv2.imencode(ext=file_type, img=image)[1].tofile(save_path)

def get_data_str():
    now = datetime.now()
    formatted_now = now.strftime("%Y-%m-%d-%H-%M-%S")
    return formatted_now

def get_past_dates(n):
    today = datetime.today()  # 获取今天的日期
    dates = [(today - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(0, n+0)]
    return dates

def open_floader(path, file_path=None):
    path = Path(path)
    current_os = platform.system()
    if current_os == "Windows":
        # 使用 explorer.exe /select, 参数选中指定文件
        if file_path:
            # file_path = Path(file_path).resolve()  # 转换为 Path 对象
            file_path = os.path.join(path, file_path)
            os.system(f'explorer.exe /select,"{file_path}"')
        else:
            os.startfile(path)  # Windows特定的方法
    elif current_os == "Darwin":  # macOS
        os.system(f"open {os.path.abspath(path)}")
    elif current_os == "Linux":
        os.system(f"xdg-open {os.path.abspath(path)}")

def unzip_file(zip_path, extract_path):
    shutil.unpack_archive(zip_path, extract_path)

def get_all_pic_path(floader_path):
    file_paths = []
    # 遍历文件夹及子文件夹
    for root, _, files in os.walk(floader_path):
        for file in files:
            # 获取完整的文件路径
            if file != '.DS_Store':
                file_paths.append(os.path.join(root, file))
    return natsorted(file_paths)

def get_pic_info(info_file_path: str, pic_path):
    with open(info_file_path, 'r', encoding='utf-8') as reader:
        lines = reader.readlines()
    reader.close()
    info = []
    for line in lines:
        line = json.loads(line.strip())
        if line['path'] == pic_path:
            info = line['info']
            break
    return info

def add_mask(image, mask, color = (255, 255, 255)):
    image_shape = image.shape
    assert image_shape == mask.shape
    color = tuple([color[0]] * image_shape[2])

    # 将原图与反转后的掩码相结合，保持mask之外的原始内容
    original_area = cv2.bitwise_and(image, ~mask)

    # 将mask之内的区域设置为白色
    corol_background = np.full(image_shape, color, dtype=np.uint8)
    corol_background = cv2.bitwise_and(corol_background, mask)

    # 合成最终图像
    rounded_image = cv2.add(corol_background, original_area)
    return rounded_image

def get_mask(corner_radius, mask_shape):
    # 创建一个带有圆角的黑色掩码
    mask_shape = (mask_shape[1], mask_shape[0], mask_shape[2])
    mask = np.full(mask_shape, 255, dtype=np.uint8)
    rows, cols, _ = mask_shape
    # 绘制矩形区域
    mask = cv2.rectangle(mask, (corner_radius, corner_radius), (cols-corner_radius, rows-corner_radius), (0, 0, 0), thickness=-1)
    mask = cv2.rectangle(mask, (corner_radius, 0), (cols-corner_radius, corner_radius), (0, 0, 0), thickness=-1)
    mask = cv2.rectangle(mask, (corner_radius, rows-corner_radius), (cols-corner_radius, rows), (0, 0, 0), thickness=-1)
    mask = cv2.rectangle(mask, (0, corner_radius), (corner_radius, rows-corner_radius), (0, 0, 0), thickness=-1)
    mask = cv2.rectangle(mask, (cols-corner_radius, corner_radius), (cols, rows-corner_radius), (0, 0, 0), thickness=-1)
    # 圆角处理
    mask = cv2.circle(mask, (corner_radius, corner_radius), corner_radius, (0, 0, 0), thickness=-1)
    mask = cv2.circle(mask, (cols-corner_radius, corner_radius), corner_radius, (0, 0, 0), thickness=-1)
    mask = cv2.circle(mask, (corner_radius, rows-corner_radius), corner_radius, (0, 0, 0), thickness=-1)
    mask = cv2.circle(mask, (cols-corner_radius, rows-corner_radius), corner_radius, (0, 0, 0), thickness=-1)
    mask = cv2.rotate(mask, cv2.ROTATE_90_CLOCKWISE)
    return mask

def cv_to_qimage(cv_img):
    try:
        height, width, _ = cv_img.shape
    except:
        height, width = cv_img.shape
    bytes_per_line = 3 * width
    # 将 BGR 转换为 RGB（Qt 使用 RGB，OpenCV 使用 BGR）
    cv_img_rgb = cv2.cvtColor(cv_img, cv2.COLOR_BGR2RGB)
    qimage = QImage(cv_img_rgb.data, width, height, bytes_per_line, QImage.Format.Format_RGB888)
    return qimage

def cv_to_qpixmap(cv_img):
    qimage = cv_to_qimage(cv_img)
    qpixmap = QPixmap.fromImage(qimage)
    return qpixmap

def calculate_distance(point1, point2):
    x1, y1 = point1
    x2, y2 = point2
    distance = math.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2)
    return distance

def point_rotate_90(point1, point_center, mode):
    """
    计算点 B(x2, y2) 围绕点 A(x1, y1) 旋转 90 度后的坐标。

    :param x1: A 点的 x 坐标
    :param y1: A 点的 y 坐标
    :param x2: B 点的 x 坐标
    :param y2: B 点的 y 坐标
    :param mode: 旋转模式，0 表示顺时针旋转，1 表示逆时针旋转
    :return: 旋转后的点 B 的新坐标 (x2', y2')
    """
    # 平移坐标系
    x1, y1 = point_center[0], point_center[1]
    x2, y2 = point1[0], point1[1]
    x_prime = x2 - x1
    y_prime = y2 - y1

    if mode == 0:
        # 顺时针旋转 90 度
        x_double_prime = y_prime
        y_double_prime = -x_prime
    elif mode == 1:
        # 逆时针旋转 90 度
        x_double_prime = -y_prime
        y_double_prime = x_prime

    # 还原坐标系
    x2_new = x_double_prime + x1
    y2_new = y_double_prime + y1

    return [x2_new, y2_new]

def pic_rotate_angle(image, mode):
    if mode:
        image = cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
    else:
        image = cv2.rotate(image, cv2.ROTATE_90_CLOCKWISE)
    return image

def add_edge(image_or, mask_angle, edge_ipx, edge_color, add_towards):
    image = copy.deepcopy(image_or)
    shift = 1
    height, width, _ = image.shape
    if 2 in add_towards:
        draw_fan(image, (width - mask_angle, mask_angle), mask_angle-edge_ipx, mask_angle, 270, 360, edge_color, thickness=1)
    if 4 in add_towards:
        draw_fan(image, (width - mask_angle, height - mask_angle), mask_angle-edge_ipx, mask_angle, 0, 90, edge_color, thickness=1)
    if 6 in add_towards:
        draw_fan(image, (mask_angle, height - mask_angle), mask_angle-edge_ipx, mask_angle, 90, 180, edge_color, thickness=1)
    if 8 in add_towards:
        draw_fan(image, (mask_angle, mask_angle), mask_angle-edge_ipx, mask_angle, 180, 270, edge_color, thickness=1)
    if edge_ipx != 1:
        if 1 in add_towards:
            cv2.rectangle(image, (mask_angle, 0), (width - mask_angle , edge_ipx), edge_color, thickness=-1)
        if 3 in add_towards:
            cv2.rectangle(image, (width - edge_ipx - shift, mask_angle), (width - shift, height - mask_angle), edge_color, thickness=-1)
        if 5 in add_towards:
            cv2.rectangle(image, (mask_angle, height - edge_ipx - shift), (width - mask_angle, height - shift), edge_color, thickness=-1)
        if 7 in add_towards:
            cv2.rectangle(image, (0, mask_angle), (edge_ipx, height - mask_angle), edge_color, thickness=-1)
    else:
        if 1 in add_towards:
            cv2.line(image, (mask_angle , 0), (width - mask_angle , 0), edge_color, 1)
        if 3 in add_towards:
            cv2.line(image, (width - shift, mask_angle), (width - shift, height - mask_angle), edge_color, 1)
        if 5 in add_towards:
            cv2.line(image, (mask_angle, height - shift), (width - mask_angle, height - shift), edge_color, 1)
        if 7 in add_towards:
            cv2.line(image, (0, mask_angle), (0, height - mask_angle), edge_color, 1)
    return image

def draw_fan(image, center, inner_radius, outer_radius, start_angle, end_angle, color, thickness=1):
    """
    在图像上绘制一个扇形区域。

    参数:
    - image: 输入图像
    - center: 圆心坐标 (x, y)
    - inner_radius: 内半径
    - outer_radius: 外半径
    - start_angle: 起始角度（度数）
    - end_angle: 结束角度（度数）
    - color: 扇形的颜色
    - thickness: 扇形的厚度 (-1 为填充)
    """
    # 定义角度范围
    if inner_radius < 0:
        inner_radius = 0
    angles = np.linspace(start_angle, end_angle, num=100)
    # 计算内半径上的点
    inner_points = [(int(center[0] + inner_radius * np.cos(np.deg2rad(angle))),
                     int(center[1] + inner_radius * np.sin(np.deg2rad(angle))))
                    for angle in angles]
    # 计算外半径上的点
    outer_points = [(int(center[0] + outer_radius * np.cos(np.deg2rad(angle))),
                     int(center[1] + outer_radius * np.sin(np.deg2rad(angle))))
                    for angle in angles]
    # 合并外半径和内半径上的点
    points = np.array(outer_points + inner_points[::-1], dtype=np.int32)
    # 使用 fillPoly 绘制扇形区域
    cv2.fillPoly(image, [points], color)
    return image

# def calculate_brightness(image):
#     """
#     计算图片的平均亮度
#     """
#     hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
#     brightness = hsv[:, :, 2].mean()  # V 通道的均值
#     return brightness

# def adjust_brightness(image, target_brightness):
#     """
#     调整图片的亮度，使其接近目标亮度
#     """
#     current_brightness = calculate_brightness(image)
#     brightness_ratio = target_brightness / current_brightness
#     # 调整 V 通道
#     hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
#     hsv = hsv.astype(np.float32)
#     hsv[:, :, 2] = np.clip(hsv[:, :, 2] * brightness_ratio, 0, 255)
#     hsv = hsv.astype(np.uint8)
#     adjusted_image = cv2.cvtColor(hsv, cv2.COLOR_HSV2BGR)
#     return adjusted_image

def download_zip(global_config, name, files = []):
    owner = global_config['owner']
    repo = global_config['repo']
    access_token = global_config['access_token']
    ref = global_config['ref']
    if files == []:
        tip = 'zipball'
    zip_url = f'https://gitee.com/api/v5/repos/{owner}/{repo}/{tip}?access_token={access_token}&ref={ref}'
    try:
        response = requests.get(zip_url, stream=True)
        if response.status_code == 200:
            with open(f"./{name}.zip", "wb") as file:
                for chunk in response.iter_content(chunk_size=1024):
                    file.write(chunk)
            return True
        else:
            return f"下载失败，状态码：{response.status_code}\n详情：{response.text}"
    except Exception as e:
        return f"下载过程中出错：{e}"
    
def download_single_file(global_config, path_in_cloud, to_path):
    owner = global_config['owner']
    repo = global_config['repo']
    access_token = global_config['access_token']
    ref = global_config['ref']

    # 构建 API URL
    url = f"https://gitee.com/api/v5/repos/{owner}/{repo}/contents/{path_in_cloud}?access_token={access_token}&ref={ref}"

    # 请求文件内容
    response = requests.get(url)

    if response.status_code == 200:
        file_info = response.json()
        file_content = base64.b64decode(file_info['content'])  # 文件内容是 Base64 编码的
        # 保存到本地
        with open(to_path, "wb") as file:
            file.write(file_content)
        return True
    else:
        return False

#从上到下, 从左到右切割的比例，和切割的模式横向或者纵向， 0为分上下，1为分左右
def split_image(input_image_path: str, output_image_path_1, output_image_path_2, proportion, toward_mode):
    try:
        pic_format = input_image_path.rsplit('.', maxsplit=1)[-1]
        # 打开输入的图片
        with Image.open(input_image_path) as img:
            # 获取图片的宽度和高度
            width, height = img.size
            if toward_mode == 0:
                # 计算分割的高度
                mid_height = int(height * proportion) 
                # 裁剪上下两部分
                upper_half = img.crop((0, 0, width, mid_height))
                lower_half = img.crop((0, mid_height, width, height))
            elif toward_mode == 1:
                # 计算分割的高度
                mid_width = int(width * proportion) 
                # 裁剪上下两部分
                upper_half = img.crop((0, 0, mid_width, height))
                lower_half = img.crop((mid_width, 0, width, height))
            # 保存裁剪后的图片到指定路径
            upper_half.save(f'{output_image_path_1}.{pic_format}', format=upper_half.format)
            lower_half.save(f'{output_image_path_2}.{pic_format}', format=lower_half.format)
            #删除原有照片
            send2trash(input_image_path)
    except:
        pass

def recursive_update(stay_data, target_data, stay_paths):
    """
    将 stay_data 中的值更新到 target_data 中，但 stay_paths 中指定的路径保持 target_data 中的值不变。

    参数说明：
        stay_data: ruamel.yaml 读取的第一个yaml数据（通常是dict）
        target_data: ruamel.yaml 读取的第二个yaml数据（通常是dict）
        stay_paths: list of list，表示需要保持 target_data 不变的路径，例如 [(a, b, c), (x, y)]
    """
    def should_stay(path):
        """
        判断当前路径 path 是否在 stay_paths 中或其子路径中。
        只要路径 path 是 stay_paths 中某路径的前缀，则也认为该路径需要保持不变，
        以防止对 stay_paths 指定路径下的子树做替换。
        """
        for sp in stay_paths:
            # 如果当前路径是 stay_paths 中路径的前缀或相等
            if len(path) <= len(sp) and all(path[i] == sp[i] for i in range(len(path))):
                return True
        return False

    def _recursive_update(sd, td, current_path=()):
        """
        递归遍历并更新 td 中的值
        """
        if not isinstance(sd, dict) or not isinstance(td, dict):
            # 如果不是字典，无法再递归，直接替换（除非在 stay_paths 中）
            if not should_stay(current_path):
                return sd
            else:
                return td

        # 都是字典，遍历 stay_data 的所有键
        for k in sd:
            new_path = current_path + (k,)
            if k in td:
                if should_stay(new_path):
                    # 路径在 stay_paths 中，跳过替换，保持 target_data 原值
                    continue
                # 递归更新
                td[k] = _recursive_update(sd[k], td[k], new_path)
            else:
                # k 不在 target_data 中，是否添加？题目未说明，这里不添加，只处理同时存在的
                pass
        return td

    _recursive_update(stay_data, target_data)
    return target_data

#计算非线性参数
def calc_no_line_number(min_number, middle_number, max_number, k_high, count_number):
    assert max_number > middle_number > min_number
    assert max_number >= count_number >= min_number
    if count_number <= middle_number:
        return count_number * k_high
    else:
        point_y = middle_number * k_high
        return point_y + (count_number - middle_number) * (max_number - point_y) / (max_number - middle_number)
    
def find_png_by_name_fast(root_folder:str, target_name:list, formate:str):
    """
    递归使用 os.scandir 以提高遍历效率，
    查找文件名为 target_name（不带扩展名）的 PNG 文件，返回第一个匹配的绝对路径，找不到返回 None。
    """
    target_lower = [i.strip().lower() for i in target_name]

    def recursive_scan(folder):
        try:
            with os.scandir(folder) as it:
                for entry in it:
                    if entry.is_file():
                        name, ext = os.path.splitext(entry.name)
                        if ext.lower() == formate and name.lower() in target_lower:
                            return entry.path
                    elif entry.is_dir():
                        found = recursive_scan(entry.path)
                        if found:
                            return found
        except PermissionError:
            # 无权限访问的文件夹忽略
            pass
        return None

    return recursive_scan(root_folder)

def to_relative_paths(abs_paths, base_path):
    rel_paths = []
    for abs_path in abs_paths:
        rel_path = os.path.relpath(abs_path, base_path)
        rel_paths.append(rel_path)
    return rel_paths

# #获取首字母
# def get_name_initials(name):
#     # 获取每个汉字的拼音首字母，Style.FIRST_LETTER表示首字母
#     initials = pinyin(name, style=Style.FIRST_LETTER)
#     # initials 是列表的列表，例如 [['z'], ['h'], ['a']]，展开并大写
#     initials_str = ''.join([item[0].upper() for item in initials])
#     return initials_str